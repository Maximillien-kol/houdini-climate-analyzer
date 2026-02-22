"""
document_parser.py
==================
Parses Rwanda Meteorology Agency (Meteorwanda) monthly forecast documents
(PDF and DOCX) downloaded from:
    https://www.meteorwanda.gov.rw/products/monthly-forecast

Extracted structured data is returned as a pandas DataFrame whose columns
are a *subset* of the schema used in rwanda_agri_climate.csv, so the rows
can be merged directly into the training dataset.

Extracted columns
-----------------
date            YYYY-MM-DD  (first day of each dekad / month)
region          one of: Kigali | Northern | Southern | Eastern | Western
source          "document"  (to distinguish from synthetic rows)
rainfall_mm     forecast rainfall in mm
temperature_c   average / mean temperature in °C
temp_min_c      forecast minimum temperature
temp_max_c      forecast maximum temperature
humidity_pct    if present
drought_index   derived (0=normal … 1=severe) from below/normal/above signals
rain_tomorrow   binary – derived from rainfall forecast
season          derived from month

Supports
--------
* pdfplumber  – primary PDF engine (best for tables)
* PyMuPDF     – fallback PDF engine
* python-docx – for .docx files
"""

from __future__ import annotations

import os
import re
import logging
from datetime import date, timedelta
from typing import Optional
import pandas as pd
import numpy as np

logger = logging.getLogger(__name__)

# ── Rwanda region aliases found in documents ──────────────────────────────────
REGION_ALIASES: dict[str, str] = {
    # canonical name → variants that appear in docs
    "Northern": ["northern", "north", "northern province"],
    "Southern": ["southern", "south", "southern province"],
    "Eastern":  ["eastern",  "east",  "eastern province"],
    "Western":  ["western",  "west",  "western province"],
    "Kigali":   ["kigali",   "kigali city", "city of kigali"],
}

# build a lookup: lowercased variant → canonical
_REGION_LOOKUP: dict[str, str] = {}
for canonical, variants in REGION_ALIASES.items():
    for v in variants:
        _REGION_LOOKUP[v] = canonical

# Rwanda bimodal season lookup (same as generate_data.py)
_SEASONS = {
    1: "dry", 2: "dry",
    3: "rainy_A", 4: "rainy_A", 5: "rainy_A",
    6: "dry", 7: "dry", 8: "dry",
    9: "rainy_B", 10: "rainy_B", 11: "rainy_B",
    12: "dry",
}

MONTH_NAMES = {
    "january": 1, "february": 2, "march": 3, "april": 4,
    "may": 5, "june": 6, "july": 7, "august": 8,
    "september": 9, "october": 10, "november": 11, "december": 12,
    "jan": 1, "feb": 2, "mar": 3, "apr": 4,
    "jun": 6, "jul": 7, "aug": 8,
    "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

# ── regex helpers ─────────────────────────────────────────────────────────────
_RE_YEAR       = re.compile(r"\b(20[12]\d)\b")
_RE_MONTH_WORD = re.compile(
    r"\b(january|february|march|april|may|june|july|august|"
    r"september|october|november|december|"
    r"jan|feb|mar|apr|jun|jul|aug|sep|oct|nov|dec)\b",
    re.I,
)
_RE_RAINFALL   = re.compile(r"(\d{1,4}(?:\.\d{1,2})?)\s*mm", re.I)
_RE_TEMP       = re.compile(r"(\d{1,2}(?:\.\d)?)\s*°?\s*[cC]")
_RE_BELOW      = re.compile(r"below[\s-]*(normal|average)", re.I)
_RE_ABOVE      = re.compile(r"above[\s-]*(normal|average)", re.I)
_RE_NORMAL     = re.compile(r"\bnormal\b", re.I)


# ═══════════════════════════════════════════════════════════════════════════════
# Public entry point
# ═══════════════════════════════════════════════════════════════════════════════

def parse_document(filepath: str) -> pd.DataFrame:
    """
    Parse a single PDF or DOCX file.

    Returns a DataFrame with the columns listed in the module docstring.
    Returns an empty DataFrame if parsing fails or yields no usable rows.
    """
    filepath = os.path.abspath(filepath)
    ext = os.path.splitext(filepath)[1].lower()

    logger.info(f"[DocParser] Parsing: {os.path.basename(filepath)}")

    try:
        if ext == ".pdf":
            return _parse_pdf(filepath)
        elif ext in (".docx", ".doc"):
            return _parse_docx(filepath)
        else:
            logger.warning(f"[DocParser] Unsupported extension: {ext}")
            return _empty_df()
    except Exception as exc:
        logger.error(f"[DocParser] Failed on {filepath}: {exc}")
        return _empty_df()


def parse_documents_folder(folder: str) -> pd.DataFrame:
    """
    Recursively scan *folder* for .pdf and .docx files and parse them all.
    Returns a combined DataFrame (deduplicated by date+region).
    """
    all_frames: list[pd.DataFrame] = []
    supported = (".pdf", ".docx", ".doc")

    for root, _dirs, files in os.walk(folder):
        for fname in sorted(files):
            if os.path.splitext(fname)[1].lower() in supported:
                fpath = os.path.join(root, fname)
                df = parse_document(fpath)
                if not df.empty:
                    all_frames.append(df)

    if not all_frames:
        logger.info("[DocParser] No documents found or all failed to parse.")
        return _empty_df()

    combined = pd.concat(all_frames, ignore_index=True)
    before = len(combined)
    combined = combined.drop_duplicates(subset=["date", "region"])
    logger.info(
        f"[DocParser] Combined {before} rows → {len(combined)} unique (date, region) pairs"
    )
    return combined


# ═══════════════════════════════════════════════════════════════════════════════
# PDF parsing
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_pdf(filepath: str) -> pd.DataFrame:
    """Try pdfplumber first, fall back to PyMuPDF text extraction."""
    try:
        import pdfplumber
        return _parse_pdf_pdfplumber(filepath)
    except ImportError:
        pass

    try:
        import fitz  # PyMuPDF
        return _parse_pdf_fitz(filepath, fitz)
    except ImportError:
        logger.error(
            "[DocParser] Neither pdfplumber nor PyMuPDF is installed. "
            "Run: pip install pdfplumber   or   pip install pymupdf"
        )
        return _empty_df()


def _parse_pdf_pdfplumber(filepath: str) -> pd.DataFrame:
    import pdfplumber

    all_text = []
    all_tables = []

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            all_text.append(txt)
            tables = page.extract_tables()
            if tables:
                all_tables.extend(tables)

    full_text = "\n".join(all_text)
    year, month = _extract_year_month(full_text, filepath)

    rows = []

    # ── try to extract from tables first ─────────────────────────────────────
    if all_tables:
        rows = _parse_tables(all_tables, year, month)

    # ── fall back to text extraction ──────────────────────────────────────────
    if not rows:
        rows = _parse_text_blocks(full_text, year, month)

    return _rows_to_df(rows)


def _parse_pdf_fitz(filepath: str, fitz) -> pd.DataFrame:
    doc = fitz.open(filepath)
    full_text = "\n".join(page.get_text() for page in doc)
    year, month = _extract_year_month(full_text, filepath)
    rows = _parse_text_blocks(full_text, year, month)
    return _rows_to_df(rows)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX parsing
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_docx(filepath: str) -> pd.DataFrame:
    try:
        from docx import Document
    except ImportError:
        logger.error(
            "[DocParser] python-docx is not installed. "
            "Run: pip install python-docx"
        )
        return _empty_df()

    doc = Document(filepath)

    # Collect full text
    para_texts = [p.text for p in doc.paragraphs]
    full_text = "\n".join(para_texts)

    year, month = _extract_year_month(full_text, filepath)

    rows = []

    # ── parse tables ──────────────────────────────────────────────────────────
    raw_tables = []
    for table in doc.tables:
        raw = []
        for row in table.rows:
            raw.append([cell.text.strip() for cell in row.cells])
        raw_tables.append(raw)

    if raw_tables:
        rows = _parse_tables(raw_tables, year, month)

    # ── fall back to paragraph scan ───────────────────────────────────────────
    if not rows:
        rows = _parse_text_blocks(full_text, year, month)

    return _rows_to_df(rows)


# ═══════════════════════════════════════════════════════════════════════════════
# Table parser (shared between PDF and DOCX)
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_tables(
    raw_tables: list[list[list[str]]],
    year: int,
    month: int,
) -> list[dict]:
    """
    Try to interpret raw cell data as regional climate rows.
    Looks for:
      - A header row containing region names
      - Data rows with numeric rainfall / temperature values
    """
    rows: list[dict] = []

    for table in raw_tables:
        if not table or len(table) < 2:
            continue

        header = [str(c).lower().strip() for c in table[0]]

        # detect which column indices map to regions
        region_cols: dict[int, str] = {}
        for idx, cell in enumerate(header):
            canonical = _REGION_LOOKUP.get(cell.strip())
            if canonical:
                region_cols[idx] = canonical

        if not region_cols:
            # maybe regions are in the first column (rows orientation)
            for row in table[1:]:
                if not row:
                    continue
                cell0 = str(row[0]).lower().strip()
                canonical = _REGION_LOOKUP.get(cell0)
                if canonical:
                    rainfall, temp_avg, drought = _extract_numerics_from_cells(
                        row[1:]
                    )
                    if rainfall is not None or temp_avg is not None:
                        rows.append(
                            _build_row(year, month, canonical, rainfall, temp_avg, None, None, drought)
                        )
            continue

        # regions are in header columns → each data row is a dekad / variable
        for row in table[1:]:
            row_label = str(row[0]).lower() if row else ""
            is_rain_row = any(k in row_label for k in ("rain", "precip", "mm"))
            is_temp_row = any(k in row_label for k in ("temp", "°", "max", "min", "mean", "avg"))

            for col_idx, canonical in region_cols.items():
                if col_idx >= len(row):
                    continue
                cell_val = str(row[col_idx]).strip()
                nums = _RE_RAINFALL.findall(cell_val) or _RE_TEMP.findall(cell_val)
                val = float(nums[0]) if nums else _try_float(cell_val)
                if val is None:
                    continue

                rainfall = val if is_rain_row else None
                temp_avg = val if is_temp_row else None
                rows.append(
                    _build_row(year, month, canonical, rainfall, temp_avg, None, None, None)
                )

    return rows


def _extract_numerics_from_cells(cells: list) -> tuple[Optional[float], Optional[float], Optional[float]]:
    """Extract rainfall(mm), temperature and drought_index from a list of cells."""
    rainfall = temp_avg = drought = None
    for cell in cells:
        text = str(cell)
        mm = _RE_RAINFALL.findall(text)
        tc = _RE_TEMP.findall(text)
        if mm and rainfall is None:
            rainfall = float(mm[0])
        if tc and temp_avg is None:
            temp_avg = float(tc[0])
        if _RE_BELOW.search(text):
            drought = 0.65
        elif _RE_ABOVE.search(text):
            drought = 0.1
        elif _RE_NORMAL.search(text):
            drought = 0.3
    return rainfall, temp_avg, drought


# ═══════════════════════════════════════════════════════════════════════════════
# Text-block fallback parser
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_text_blocks(text: str, year: int, month: int) -> list[dict]:
    """
    Line-by-line scan.  When a line mentions a region, look for numeric
    rainfall / temperature values in the surrounding ±3 lines.
    """
    lines = text.splitlines()
    rows: list[dict] = []
    used_regions: set[str] = set()

    for i, line in enumerate(lines):
        lower = line.lower()

        # find canonical region
        canonical = None
        for alias, canon in sorted(_REGION_LOOKUP.items(), key=lambda x: -len(x[0])):
            if alias in lower:
                canonical = canon
                break

        if canonical is None:
            continue

        # look in surrounding lines for numeric values
        context = "\n".join(lines[max(0, i - 3): min(len(lines), i + 4)])
        rainfall_matches = _RE_RAINFALL.findall(context)
        temp_matches = _RE_TEMP.findall(context)

        rainfall = float(rainfall_matches[0]) if rainfall_matches else None
        temp_avg = float(temp_matches[0]) if temp_matches else None

        # drought signal from keywords
        drought = None
        if _RE_BELOW.search(context):
            drought = 0.65
        elif _RE_ABOVE.search(context):
            drought = 0.1
        elif _RE_NORMAL.search(context):
            drought = 0.3

        # avoid duplication for the same region in one document
        key = (canonical, i // 5)  # roughly group by position
        if key not in used_regions and (rainfall is not None or temp_avg is not None):
            used_regions.add(key)
            rows.append(
                _build_row(year, month, canonical, rainfall, temp_avg, None, None, drought)
            )

    return rows


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _build_row(
    year: int,
    month: int,
    region: str,
    rainfall_mm: Optional[float],
    temperature_c: Optional[float],
    temp_min_c: Optional[float],
    temp_max_c: Optional[float],
    drought_index: Optional[float],
) -> dict:
    """Create one structured row for the given month/region."""
    d = date(year, month, 1)

    # Derive drought index from rainfall if not already set
    if drought_index is None:
        if rainfall_mm is not None:
            # Compare against Rwanda's seasonal norms (rough baseline)
            baseline = {
                1: 50, 2: 70, 3: 130, 4: 160, 5: 120,
                6: 30, 7: 15, 8: 20, 9: 80, 10: 110, 11: 120, 12: 60,
            }.get(month, 80)
            ratio = rainfall_mm / (baseline + 1)
            drought_index = float(np.clip(1 - ratio + 0.1, 0, 1))
        else:
            drought_index = 0.3  # assume normal

    # Derive rain_tomorrow (binary) from rainfall
    rain_tomorrow = 1 if (rainfall_mm or 0) > 20 else 0

    return {
        "date":          d.strftime("%Y-%m-%d"),
        "region":        region,
        "season":        _SEASONS.get(month, "dry"),
        "source":        "document",
        "rainfall_mm":   round(float(rainfall_mm), 2) if rainfall_mm is not None else np.nan,
        "temperature_c": round(float(temperature_c), 2) if temperature_c is not None else np.nan,
        "temp_min_c":    round(float(temp_min_c), 2) if temp_min_c is not None else np.nan,
        "temp_max_c":    round(float(temp_max_c), 2) if temp_max_c is not None else np.nan,
        "drought_index": round(float(drought_index), 4),
        "rain_tomorrow": rain_tomorrow,
    }


def _extract_year_month(text: str, filepath: str) -> tuple[int, int]:
    """
    Try to determine the forecast year and month from document text or filename.
    Falls back to (current_year, current_month) with a warning.

    Supported filename patterns (case-insensitive):
      Weather_Forecast_for_December_2025.pdf   ← primary Meteorwanda format
      march_2023.pdf
      2023_march.pdf
      2023-03.pdf / 2023_03.pdf
    """
    # 1. Try filename – works for any pattern that contains a month word and a
    #    4-digit year anywhere in the name, e.g. Weather_Forecast_for_May_2024.pdf
    #    Underscores are replaced with spaces so \b word-boundaries fire correctly
    #    across patterns like  ..._December_2025.pdf
    fname_raw = os.path.basename(filepath).lower()
    fname     = fname_raw.replace("_", " ").replace("-", " ")
    year = month = None

    fname_year = _RE_YEAR.search(fname)
    if fname_year:
        year = int(fname_year.group(1))

    fname_month_word = _RE_MONTH_WORD.search(fname)
    if fname_month_word:
        month = MONTH_NAMES[fname_month_word.group(1).lower()]

    # numeric month still uses the raw name to catch  2023-03  or  2023_03
    fname_month_num = re.search(r"[-_](\d{2})[-_.]", fname_raw)
    if month is None and fname_month_num:
        m = int(fname_month_num.group(1))
        if 1 <= m <= 12:
            month = m

    # Also check parent folder – if structured as monthly_forecasts/2023/
    parent = os.path.basename(os.path.dirname(filepath))
    if year is None and _RE_YEAR.match(parent):
        year = int(parent)

    # 2. Scan document text for year / month
    if year is None:
        m_year = _RE_YEAR.search(text)
        if m_year:
            year = int(m_year.group(1))

    if month is None:
        m_mon = _RE_MONTH_WORD.search(text)
        if m_mon:
            month = MONTH_NAMES[m_mon.group(1).lower()]

    # 3. Defaults
    today = date.today()
    if year is None:
        logger.warning(f"[DocParser] Could not determine year for {fname}, using {today.year}")
        year = today.year
    if month is None:
        logger.warning(f"[DocParser] Could not determine month for {fname}, using {today.month}")
        month = today.month

    return year, month


def _rows_to_df(rows: list[dict]) -> pd.DataFrame:
    if not rows:
        return _empty_df()
    df = pd.DataFrame(rows)
    df["date"] = pd.to_datetime(df["date"])
    return df


def _empty_df() -> pd.DataFrame:
    return pd.DataFrame(columns=[
        "date", "region", "season", "source",
        "rainfall_mm", "temperature_c", "temp_min_c", "temp_max_c",
        "drought_index", "rain_tomorrow",
    ])


def _try_float(val: str) -> Optional[float]:
    try:
        return float(val.replace(",", "."))
    except (ValueError, AttributeError):
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# CLI convenience
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) < 2:
        print("Usage: python document_parser.py  <file_or_folder>")
        sys.exit(1)
    target = sys.argv[1]
    if os.path.isdir(target):
        out = parse_documents_folder(target)
    else:
        out = parse_document(target)
    print(out.to_string())
    print(f"\nShape: {out.shape}")
